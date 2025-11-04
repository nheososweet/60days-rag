"""
Document Upload & Embedding API
Handles file upload, text extraction, and embedding

LEARNING NOTES:
===============
FastAPI endpoints for:
1. Upload file (PDF/DOCX)
2. Extract text
3. Embed & store in vector DB
4. List documents
5. Delete document
"""

from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import List, Optional
import uuid
import os
import PyPDF2
import docx
import io

from app.services.embedding_service import EmbeddingService
from app.services.vector_db_service import VectorDBService

router = APIRouter(prefix="/api/documents", tags=["documents"])

# Initialize services
embedding_service = EmbeddingService()
vector_db = VectorDBService()

# Store uploaded files temporarily
UPLOAD_DIR = "./uploaded_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)


class EmbedRequest(BaseModel):
    """Request body for embedding endpoint"""
    document_id: str
    filename: str


@router.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """
    Upload PDF or DOCX file
    
    LEARNING - FILE UPLOAD:
    ======================
    UploadFile: FastAPI class for handling file uploads
    - Automatically handles multipart/form-data
    - Provides file metadata (filename, content_type, size)
    - Streams file content (efficient for large files)
    
    Steps:
    1. Validate file type
    2. Generate unique ID
    3. Save file temporarily
    4. Return document_id for next step (embedding)
    
    Returns:
        JSON with document_id and filename
    """
    # Validate file type
    # LEARNING: MIME types for different file formats
    allowed_types = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type. Allowed: PDF, DOCX. Got: {file.content_type}"
        )
    
    # Generate unique document ID
    # LEARNING: UUID = Universally Unique Identifier
    document_id = str(uuid.uuid4())
    
    # Save file
    file_path = os.path.join(UPLOAD_DIR, f"{document_id}_{file.filename}")
    
    try:
        # Read and save file content
        # LEARNING: UploadFile.read() returns bytes
        content = await file.read()
        
        with open(file_path, "wb") as f:
            f.write(content)
        
        print(f"📁 Saved file: {file_path}")
        
        return {
            "success": True,
            "document_id": document_id,
            "filename": file.filename,
            "size": len(content),
            "message": "File uploaded successfully"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file
    
    LEARNING - PDF PARSING:
    ======================
    PyPDF2: Library to read PDF files
    - PDFs store text + formatting + images
    - Need to extract just the text content
    - Some PDFs are scanned images → need OCR (not covered here)
    
    Process:
    1. Open PDF with PyPDF2
    2. Loop through pages
    3. Extract text from each page
    4. Combine all text
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text as string
    """
    try:
        text = ""
        
        # Open PDF
        with open(file_path, 'rb') as f:
            # Create PDF reader
            pdf_reader = PyPDF2.PdfReader(f)
            
            # Get number of pages
            num_pages = len(pdf_reader.pages)
            print(f"📄 PDF has {num_pages} pages")
            
            # Extract text from each page
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                text += page_text + "\n\n"
                print(f"  ✓ Extracted page {page_num + 1}")
        
        return text.strip()
        
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file
    
    LEARNING - DOCX PARSING:
    =======================
    python-docx: Library to read Word documents
    - DOCX is actually a ZIP file with XML inside
    - python-docx handles the complexity
    - Extracts paragraphs and tables
    
    Process:
    1. Open DOCX with python-docx
    2. Loop through paragraphs
    3. Extract text from each paragraph
    4. Combine all text
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text as string
    """
    try:
        # Open DOCX
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        text = ""
        for i, paragraph in enumerate(doc.paragraphs):
            text += paragraph.text + "\n"
        
        print(f"📄 DOCX has {len(doc.paragraphs)} paragraphs")
        
        return text.strip()
        
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


@router.post("/embed")
async def embed_document(request: EmbedRequest):
    """
    Process document: extract text, embed, store in vector DB
    
    LEARNING - COMPLETE PIPELINE:
    ============================
    This is the main processing endpoint
    Combines all services:
    1. File → Text (extraction)
    2. Text → Chunks (splitting)
    3. Chunks → Embeddings (Gemini API)
    4. Embeddings → Vector DB (ChromaDB)
    
    Steps:
    1. Find uploaded file
    2. Extract text based on file type
    3. Chunk text
    4. Embed chunks with Gemini
    5. Store in ChromaDB
    6. Clean up temp file
    
    Args:
        request: EmbedRequest with document_id and filename
        
    Returns:
        JSON with processing stats
    """
    document_id = request.document_id
    filename = request.filename
    
    # Find file
    file_path = os.path.join(UPLOAD_DIR, f"{document_id}_{filename}")
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    try:
        print(f"\n{'='*60}")
        print(f"📚 PROCESSING DOCUMENT: {filename}")
        print(f"{'='*60}\n")
        
        # Step 1: Extract text
        print("📄 Step 1: Extracting text...")
        
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        print(f"✓ Extracted {len(text)} characters")
        
        # Step 2: Embed document
        print("\n🎯 Step 2: Creating embeddings...")
        embeddings_data = embedding_service.process_document(text)
        
        # Step 3: Store in vector DB
        print("\n💾 Step 3: Storing in vector database...")
        result = vector_db.add_document(
            document_id=document_id,
            embeddings_data=embeddings_data,
            metadata={
                "filename": filename,
                "file_path": file_path,
                "text_length": len(text)
            }
        )
        
        # Step 4: Clean up (optional - or keep for re-processing)
        # os.remove(file_path)
        # print(f"\n🗑️  Cleaned up temp file")
        
        print(f"\n{'='*60}")
        print(f"✅ COMPLETE! Document ready for RAG")
        print(f"{'='*60}\n")
        
        return {
            "success": True,
            "document_id": document_id,
            "filename": filename,
            "chunks_count": result['chunks_stored'],
            "text_length": len(text),
            "message": "Document embedded and stored successfully"
        }
        
    except Exception as e:
        print(f"❌ Processing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")


@router.get("/stats")
async def get_stats():
    """
    Get vector database statistics
    
    Returns:
        JSON with DB stats
    """
    try:
        stats = vector_db.get_collection_stats()
        return {
            "success": True,
            **stats
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/list")
async def list_documents():
    """
    Lấy danh sách tất cả documents đã embedding (List all embedded documents)
    
    =============================================================================
    LEARNING - LIST DOCUMENTS API
    =============================================================================
    
    Endpoint này:
    1. Query vector DB để lấy all documents
    2. Group chunks theo document_id
    3. Aggregate metadata (filename, chunks count, etc.)
    4. Return formatted list cho UI
    
    Response structure:
    {
        "success": true,
        "count": 3,
        "documents": [
            {
                "document_id": "uuid-123",
                "filename": "report.pdf",
                "chunks_count": 5,
                "total_words": 2500,
                "metadata": {...}
            },
            ...
        ]
    }
    
    Use case: 
    - UI load page → call API này
    - Display documents trong table/list
    - User xem documents đã upload
    
    Returns:
        JSON with list of documents
    """
    try:
        documents = vector_db.list_all_documents()
        
        return {
            "success": True,
            "count": len(documents),
            "documents": documents
        }
    except Exception as e:
        print(f"❌ List documents error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{document_id}")
async def get_document_details(document_id: str):
    """
    Lấy chi tiết 1 document cụ thể (Get specific document details)
    
    =============================================================================
    LEARNING - DOCUMENT DETAILS API
    =============================================================================
    
    Endpoint này:
    1. Query document by ID
    2. Return ALL chunks với full details
    3. Include text previews
    
    Response structure:
    {
        "success": true,
        "document": {
            "document_id": "uuid-123",
            "filename": "report.pdf",
            "chunks_count": 5,
            "chunks": [
                {
                    "chunk_id": "uuid-123::chunk_0",
                    "chunk_index": 0,
                    "text": "Full chunk text...",
                    "text_preview": "First 200 chars...",
                    "words": 500,
                    "length": 3000
                },
                ...
            ]
        }
    }
    
    Use case:
    - User clicks "View Details" button
    - Show modal/page with chunk previews
    - Debug: Check if chunking worked correctly
    
    Args:
        document_id: Document ID to retrieve
    
    Returns:
        JSON with document details
    """
    try:
        document = vector_db.get_document_by_id(document_id)
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        return {
            "success": True,
            "document": document
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Get document error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/{document_id}")
async def delete_document(document_id: str):
    """
    Xóa document khỏi vector DB (Delete document from vector DB)
    
    =============================================================================
    LEARNING - DELETE DOCUMENT API
    =============================================================================
    
    Endpoint này:
    1. Find all chunks của document
    2. Delete all chunks from ChromaDB
    3. Optional: Delete file from disk (nếu muốn)
    
    Process:
    - User clicks "Delete" button trong UI
    - UI call DELETE /api/documents/{document_id}
    - Backend xóa all chunks khỏi vector DB
    - UI refresh list để remove document khỏi display
    
    IMPORTANT: 
    - Sau khi xóa, document KHÔNG thể recover!
    - Consider adding "soft delete" cho production
    - Có thể keep file on disk nhưng xóa embeddings
    
    Args:
        document_id: Document to delete
        
    Returns:
        JSON with deletion status
    """
    try:
        result = vector_db.delete_document(document_id)
        
        if result.get('success'):
            # Optional: Also delete file from disk
            # file_path = ...
            # if os.path.exists(file_path):
            #     os.remove(file_path)
            
            return {
                "success": True,
                "message": f"Deleted {result['chunks_deleted']} chunks",
                "chunks_deleted": result['chunks_deleted']
            }
        else:
            raise HTTPException(status_code=404, detail="Document not found")
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Delete document error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# LEARNING - HOW TO TEST:
# ======================
"""
1. Start FastAPI:
   cd 60days-rag
   uvicorn main:app --reload

2. Upload file:
   POST http://localhost:8000/api/documents/upload
   Body: form-data with "file" field
   
3. Embed document:
   POST http://localhost:8000/api/documents/embed
   Body: {
       "document_id": "uuid-from-upload",
       "filename": "your-file.pdf"
   }

4. Check stats:
   GET http://localhost:8000/api/documents/stats

5. Now ready for RAG queries!
"""
